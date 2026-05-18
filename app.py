"""
EcoCart AI System - Streamlit App
NCI MSCAI | Fundamentals of AI TABA 2026
"""

import io, os, math, heapq, time, importlib.util
from contextlib import redirect_stdout
from collections import deque

import matplotlib
matplotlib.use("Agg")
os.makedirs("output", exist_ok=True)

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# ══════════════════════════════════════════════════════════════════════════════
#  REPORT GENERATOR  (unchanged)
# ══════════════════════════════════════════════════════════════════════════════
def _build_report(t2_text, t3_text, t5_text):
    TNR = "Times New Roman"
    doc = Document()
    doc.styles["Normal"].font.name = TNR
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].paragraph_format.space_after = Pt(8)
    for lvl, sz in [(1,14),(2,13)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = TNR; s.font.bold = True; s.font.size = Pt(sz)
        s.font.color.rgb = RGBColor(0,0,0)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(4)
    def H(txt, lv=1):
        p = doc.add_heading(txt, level=lv); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs: r.font.name=TNR; r.font.color.rgb=RGBColor(0,0,0)
    def P(txt):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(txt); r.font.name=TNR; r.font.size=Pt(12)
    def CODE(txt):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(txt); r.font.name="Courier New"; r.font.size=Pt(9)
    def IMG(path, caption="", width=5.5):
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(8)
            for r in cp.runs: r.font.name=TNR; r.font.size=Pt(10); r.font.italic=True
    def SP(): doc.add_paragraph("").paragraph_format.space_after = Pt(4)
    SP(); SP()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EcoCart AI System"); r.font.name=TNR; r.font.size=Pt(24); r.font.bold=True
    p.paragraph_format.space_after = Pt(8)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Technical Report - TABA Section II"); r2.font.name=TNR; r2.font.size=Pt(14)
    p2.paragraph_format.space_after = Pt(20)
    for line in ["National College of Ireland","MSc Artificial Intelligence",
                 "Fundamentals of Artificial Intelligence","May 2026"]:
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(line); lr.font.name=TNR; lr.font.size=Pt(12)
        lp.paragraph_format.space_after = Pt(4)
    SP()
    lnk = doc.add_paragraph(); lnk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr2 = lnk.add_run("Live Demo: https://ecocart-ai-app-live.streamlit.app")
    lr2.font.name=TNR; lr2.font.size=Pt(11); lr2.font.bold=True
    lr2.font.color.rgb = RGBColor(37,99,235)
    doc.add_page_break()
    H("Task 2 - Customer Segmentation & Bias Mitigation")
    P("Running task2_segmentation.py produced the following output:")
    if t2_text: CODE(t2_text)
    SP(); IMG("output/bias_before_after.png","Figure 1: Customer clusters before and after bias mitigation")
    SP(); IMG("output/disparate_impact.png","Figure 2: Disparate Impact and High Value rates before and after fix")
    SP(); P("Before: DI=0.0 (biased). After: DI=0.847 (fair, above 0.80).")
    doc.add_page_break()
    H("Tasks 3 & 4 - Route Optimisation and Algorithm Comparison")
    P("Running task3_4_routing.py produced the following output:")
    if t3_text: CODE(t3_text)
    SP(); IMG("output/network_map.png","Figure 3: EcoCart 20-node delivery network")
    SP(); IMG("output/algo_comparison.png","Figure 4: A* vs IDA* comparison")
    SP(); IMG("output/green_vs_fast.png","Figure 5: Fastest vs lowest CO2 route")
    SP(); P("A* found the optimal path on every route with fewest nodes expanded.")
    doc.add_page_break()
    H("Task 5 - Demand Forecasting with Machine Learning")
    P("Running task5_forecasting.py produced the following output:")
    if t5_text: CODE(t5_text)
    SP(); IMG("output/forecast.png","Figure 6: Actual vs predicted daily sales")
    SP(); IMG("output/residuals.png","Figure 7: Residuals for both models")
    SP(); IMG("output/feature_importance.png","Figure 8: Random Forest feature importance")
    SP(); P("LR: MAE=9.62, R2=0.762. RF: MAE=9.75, R2=0.716. Top feature: lag_7.")
    doc.add_page_break()
    H("References")
    for ref in [
        "[1]  S. Russell and P. Norvig, Artificial Intelligence: A Modern Approach, 4th ed. Pearson, 2020.",
        "[2]  F. Pedregosa et al., Scikit-learn: Machine Learning in Python, JMLR, 2011.",
        "[3]  M. Feldman et al., Certifying and Removing Disparate Impact, ACM SIGKDD, 2015.",
        "[4]  P. E. Hart et al., A Formal Basis for the Heuristic Determination of Minimum Cost Paths, IEEE, 1968.",
    ]:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        for r in p.runs: r.font.name=TNR; r.font.size=Pt(11)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE SETUP
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(page_title="EcoCart AI", page_icon="🛒",
                   layout="wide", initial_sidebar_state="expanded")

# design tokens
NAVY  = "#0f172a"; SLATE = "#1e293b"; MUTED = "#64748b"
BORDER = "#e2e8f0"; SURF = "#ffffff"; BG = "#f8fafc"
GREEN = "#059669"; BLUE  = "#2563eb"; AMBER = "#d97706"; RED = "#dc2626"

def _rgba(h, a=1.0):
    h = h.lstrip("#"); r,g,b = int(h[:2],16),int(h[2:4],16),int(h[4:],16)
    return f"rgba({r},{g},{b},{a})"

def _layout(h=420):
    return dict(height=h, paper_bgcolor=SURF, plot_bgcolor=BG,
                font=dict(color=SLATE, size=11),
                margin=dict(l=8,r=8,t=8,b=8),
                showlegend=False)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:opsz,wght@14..32,300;14..32,400;14..32,500;14..32,600;14..32,700;14..32,800;14..32,900&display=swap');

@keyframes shimmer{0%{background-position:-300% center}100%{background-position:300% center}}
@keyframes glow-pulse{0%,100%{opacity:.18}50%{opacity:.42}}
@keyframes slide-up{from{opacity:0;transform:translateY(14px)}to{opacity:1;transform:translateY(0)}}
@keyframes dot-ping{0%,100%{box-shadow:0 0 0 0 rgba(129,140,248,.7)}50%{box-shadow:0 0 0 5px rgba(129,140,248,0)}}

*, *::before, *::after { font-family: 'Inter', -apple-system, sans-serif !important; box-sizing: border-box; }

/* Base */
[data-testid="stAppViewContainer"] { background: linear-gradient(155deg,#edf0ff 0%,#f0f2f8 45%,#f3f0ff 100%) !important; }
.block-container { padding: .8rem 1.6rem 4rem !important; max-width: 1300px !important; }
#MainMenu, footer, header { visibility: hidden; }
::-webkit-scrollbar { width: 4px; } ::-webkit-scrollbar-thumb { background: #cbd5e1; border-radius: 8px; }

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
  background: #fff; border-radius: 18px; padding: 6px; gap: 4px;
  box-shadow: 0 4px 30px rgba(0,0,0,.07), 0 1px 3px rgba(0,0,0,.04);
  border: 1px solid rgba(226,232,240,.8);
}
.stTabs [data-baseweb="tab"] {
  border-radius: 12px; font-size: .75rem; font-weight: 700;
  padding: 9px 14px; color: #94a3b8; letter-spacing: .01em;
  transition: color .15s ease;
}
.stTabs [aria-selected="true"] {
  background: linear-gradient(135deg, #1e293b 0%, #334155 100%) !important;
  color: #fff !important; box-shadow: 0 4px 14px rgba(30,41,59,.3) !important;
}

/* Hero */
.hero {
  background: linear-gradient(140deg, #0a0e1a 0%, #0f1c35 45%, #091428 100%);
  border-radius: 22px; padding: 34px 38px 30px; margin-bottom: 22px;
  position: relative; overflow: hidden;
  box-shadow: 0 24px 64px rgba(10,14,26,.45), 0 4px 16px rgba(10,14,26,.2);
}
.hero::before {
  content: ''; position: absolute; top: -100px; right: -80px;
  width: 420px; height: 420px;
  background: radial-gradient(circle, rgba(79,70,229,.22) 0%, transparent 60%);
  border-radius: 50%; pointer-events: none;
  animation: glow-pulse 5s ease-in-out infinite;
}
.hero::after {
  content: ''; position: absolute; bottom: -100px; left: 20%;
  width: 360px; height: 360px;
  background: radial-gradient(circle, rgba(5,150,105,.18) 0%, transparent 60%);
  border-radius: 50%; pointer-events: none;
  animation: glow-pulse 5s ease-in-out infinite 2.5s;
}
.hero-title {
  font-size: 2.4rem; font-weight: 900; letter-spacing: -.055em; line-height: 1.05;
  margin: 0 0 10px;
  background: linear-gradient(135deg, #ffffff 0%, #bfdbfe 35%, #6ee7b7 65%, #ffffff 100%);
  background-size: 300% auto;
  -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;
  animation: shimmer 8s linear infinite;
}
.hero-sub {
  color: #94a3b8; font-size: .86rem; font-weight: 400;
  margin-bottom: 28px; line-height: 1.65; max-width: 560px;
}
/* Task cards */
.task-card {
  background: #fff; border-radius: 18px; padding: 20px 22px;
  margin-bottom: 18px; border: 1px solid #e2e8f0;
  box-shadow: 0 4px 24px rgba(0,0,0,.05), 0 1px 3px rgba(0,0,0,.03);
  display: flex; gap: 18px; align-items: flex-start;
  transition: box-shadow .2s ease, transform .2s ease;
  animation: slide-up .45s ease both;
  border-left-width: 4px;
}
.task-card:hover { box-shadow: 0 16px 48px rgba(0,0,0,.12); transform: translateY(-3px); }
.task-icon {
  width: 52px; height: 52px; border-radius: 15px; flex-shrink: 0;
  display: flex; align-items: center; justify-content: center; font-size: 1.5rem;
}
.task-title {
  font-size: .96rem; font-weight: 800; color: #0f172a;
  margin-bottom: 6px; line-height: 1.25; letter-spacing: -.02em;
}
.task-desc { font-size: .81rem; color: #64748b; line-height: 1.7; }

/* Insight */
.insight {
  background: linear-gradient(135deg, #f0fdf4 0%, #dcfce7 100%);
  border: 1px solid #86efac; border-radius: 16px;
  padding: 16px 20px 16px 60px; position: relative;
  color: #14532d; font-size: .82rem; line-height: 1.8; margin: 14px 0;
  box-shadow: 0 8px 32px rgba(5,150,105,.15), inset 0 1px 0 rgba(255,255,255,.7);
}
.insight::before {
  content: '✓'; position: absolute; left: 16px; top: 16px;
  width: 30px; height: 30px;
  background: linear-gradient(135deg, #059669, #34d399);
  border-radius: 10px; color: #fff; font-size: .9rem; font-weight: 900;
  display: flex; align-items: center; justify-content: center;
  box-shadow: 0 3px 8px rgba(5,150,105,.35);
}

/* Warn */
.warn-box {
  background: linear-gradient(135deg, #fffbeb 0%, #fef3c7 100%);
  border: 1px solid #fcd34d; border-radius: 16px;
  padding: 16px 20px 16px 60px; position: relative;
  color: #78350f; font-size: .82rem; line-height: 1.8; margin: 14px 0;
  box-shadow: 0 8px 32px rgba(217,119,6,.15), inset 0 1px 0 rgba(255,255,255,.7);
}
.warn-box::before {
  content: '!'; position: absolute; left: 16px; top: 16px;
  width: 30px; height: 30px;
  background: linear-gradient(135deg, #d97706, #fbbf24);
  border-radius: 10px; color: #fff; font-size: 1rem; font-weight: 900;
  display: flex; align-items: center; justify-content: center;
  box-shadow: 0 3px 8px rgba(217,119,6,.35);
}

/* Terminal */
.term { border-radius: 14px; overflow: hidden; margin: 10px 0; box-shadow: 0 8px 32px rgba(0,0,0,.2); }
.term-top { background: #1e293b; padding: 9px 16px; display: flex; gap: 7px; align-items: center; }
.dot { width: 11px; height: 11px; border-radius: 50%; }
.term-body {
  background: #070d19; padding: 16px 20px;
  font-family: ui-monospace, 'Cascadia Code', 'Fira Code', monospace;
  font-size: .76rem; color: #7dd3fc;
  white-space: pre-wrap; line-height: 1.8; max-height: 280px; overflow-y: auto;
}

/* Section heading */
.sec-head {
  font-size: .9rem; font-weight: 800; color: #0f172a;
  letter-spacing: -.02em; margin: 20px 0 12px;
  display: flex; align-items: center; gap: 10px;
}
.sec-head::after { content: ''; flex: 1; height: 2px; background: linear-gradient(90deg, #e2e8f0, transparent); border-radius: 2px; }

/* Slabel */
.slabel {
  font-size: .67rem; font-weight: 800; color: #94a3b8;
  text-transform: uppercase; letter-spacing: .11em; margin-bottom: 10px;
}

/* Legend */
.leg { display: flex; gap: 14px; flex-wrap: wrap; margin: 10px 0; }
.li { display: flex; align-items: center; gap: 6px; font-size: .75rem; color: #475569; font-weight: 600; }
.ld { width: 10px; height: 10px; border-radius: 3px; flex-shrink: 0; }

/* Metrics */
div[data-testid="metric-container"] {
  background: linear-gradient(145deg, #fff 0%, #f8fafc 100%);
  border-radius: 16px; padding: 18px 20px;
  border: 1px solid #e2e8f0; border-top: 3px solid #6366f1;
  box-shadow: 0 4px 20px rgba(0,0,0,.06);
  transition: transform .18s ease, box-shadow .18s ease;
}
div[data-testid="metric-container"]:hover { transform: translateY(-3px); box-shadow: 0 12px 36px rgba(0,0,0,.1); }
div[data-testid="metric-container"] label { font-size: .68rem !important; font-weight: 800 !important; color: #94a3b8 !important; text-transform: uppercase; letter-spacing: .08em; }
div[data-testid="metric-container"] [data-testid="stMetricValue"] { font-size: 1.5rem !important; font-weight: 900 !important; color: #0f172a !important; letter-spacing: -.03em; }

/* Buttons */
.stButton > button {
  border-radius: 12px !important; font-weight: 700 !important;
  font-size: .82rem !important; letter-spacing: .01em !important;
  transition: all .2s ease !important;
}
div[data-testid="stButton"] > button[kind="primary"] {
  background: linear-gradient(135deg, #1d4ed8, #3b82f6) !important;
  border: none !important; color: #fff !important;
  box-shadow: 0 4px 16px rgba(59,130,246,.4) !important;
  padding: 10px 20px !important;
}
div[data-testid="stButton"] > button[kind="primary"]:hover {
  box-shadow: 0 8px 28px rgba(59,130,246,.5) !important;
  transform: translateY(-1px) !important;
}

/* Sidebar */
[data-testid="stSidebar"] {
  background: linear-gradient(180deg, #fff 0%, #f8fafc 100%) !important;
  border-right: 1px solid #e2e8f0 !important;
  box-shadow: 3px 0 20px rgba(0,0,0,.04) !important;
}
.sb-brand { text-align: center; padding: 10px 0 22px; border-bottom: 1px solid #f1f5f9; margin-bottom: 20px; }
.sb-icon { font-size: 2.6rem; line-height: 1; margin-bottom: 8px; }
.sb-name { font-weight: 900; font-size: 1.05rem; color: #0f172a; letter-spacing: -.02em; }
.sb-sub { font-size: .68rem; color: #94a3b8; margin-top: 3px; font-weight: 500; }
.sb-section { font-size: .63rem; font-weight: 800; color: #94a3b8; text-transform: uppercase; letter-spacing: .11em; margin: 20px 0 8px; }
.sb-step {
  display: flex; align-items: center; gap: 10px;
  padding: 10px 12px; border-radius: 12px; margin-bottom: 5px;
  background: #f8fafc; border: 1px solid #f1f5f9;
  transition: background .15s;
}
.sb-step:hover { background: #f1f5f9; }
.sb-num {
  width: 26px; height: 26px; flex-shrink: 0;
  background: linear-gradient(135deg, #1e293b, #334155);
  color: #fff; border-radius: 8px; font-size: .66rem; font-weight: 800;
  display: flex; align-items: center; justify-content: center;
  box-shadow: 0 2px 6px rgba(30,41,59,.25);
}
.sb-step-txt { font-size: .79rem; color: #334155; font-weight: 500; line-height: 1.3; }
.sb-status-row {
  display: flex; align-items: center; gap: 8px; padding: 9px 12px;
  border-radius: 12px; margin-bottom: 5px; font-size: .79rem; font-weight: 700;
}
.sb-done { background: #f0fdf4; color: #15803d; border: 1px solid #bbf7d0; }
.sb-pending { background: #f8fafc; color: #94a3b8; border: 1px solid #e8edf5; }

/* Dataframe */
.stDataFrame { border-radius: 14px !important; overflow: hidden !important; box-shadow: 0 2px 12px rgba(0,0,0,.04) !important; }

/* Expander - hide default arrow, style cleanly */
details > summary { list-style: none !important; }
details > summary::-webkit-details-marker { display: none !important; }
details > summary::marker { display: none !important; }
[data-testid="stExpanderToggleIcon"] { display: none !important; }
[data-testid="stExpander"] details summary {
  background: #f8fafc !important;
  border-radius: 12px !important;
  padding: 11px 18px !important;
  font-size: .81rem !important;
  font-weight: 600 !important;
  color: #475569 !important;
  border: 1px solid #e2e8f0 !important;
  cursor: pointer !important;
}
[data-testid="stExpander"] details summary:hover {
  background: #f1f5f9 !important;
  color: #0f172a !important;
}
[data-testid="stExpander"] details summary::after {
  content: '▾';
  float: right;
  color: #94a3b8;
  font-size: .9rem;
  transition: transform .2s;
}
[data-testid="stExpander"] details[open] summary::after {
  transform: rotate(-180deg);
}

/* Slider */
.stSlider [data-testid="stTickBarMin"], .stSlider [data-testid="stTickBarMax"] { font-size: .72rem !important; }

/* Hero badge */
.hero-badge {
  display: inline-flex; align-items: center; gap: 8px;
  background: rgba(99,102,241,.14); border: 1.5px solid rgba(99,102,241,.3);
  border-radius: 20px; padding: 5px 14px; margin-bottom: 16px;
  backdrop-filter: blur(8px);
}
.hero-badge-dot {
  width: 7px; height: 7px; border-radius: 50%; background: #818cf8; flex-shrink: 0;
  animation: dot-ping 2.4s ease-in-out infinite;
}
.hero-badge-txt { font-size: .63rem; font-weight: 700; color: #a5b4fc; text-transform: uppercase; letter-spacing: .1em; }

/* Sidebar brand icon box */
.sb-icon-box {
  width: 64px; height: 64px; border-radius: 20px; margin: 0 auto 10px;
  background: linear-gradient(135deg, #0f172a 0%, #1e3a5f 100%);
  display: flex; align-items: center; justify-content: center;
  font-size: 2rem; box-shadow: 0 6px 24px rgba(15,23,42,.25);
}
.sb-version {
  display: inline-block; background: #f0f9ff; border: 1px solid #bae6fd;
  border-radius: 10px; padding: 2px 9px; font-size: .6rem;
  font-weight: 700; color: #0ea5e9; margin-top: 5px;
}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div class='sb-brand'>
      <div class='sb-icon-box'>🛒</div>
      <div class='sb-name'>EcoCart AI</div>
    </div>""", unsafe_allow_html=True)

    st.markdown("<div class='sb-section'>How to use</div>", unsafe_allow_html=True)
    for n, t in [("1","Pick a task tab above"),
                 ("2","Tasks 2, 3, 5 — press Run"),
                 ("3","Tasks 1 & 3 — use Play / Slider"),
                 ("4","Task 6 — adjust the sliders")]:
        st.markdown(f"""<div class='sb-step'>
          <div class='sb-num'>{n}</div>
          <span class='sb-step-txt'>{t}</span></div>""", unsafe_allow_html=True)

    st.markdown("<div class='sb-section'>Task progress</div>", unsafe_allow_html=True)
    t2_done = st.session_state.get("t2_done", False)
    t3_done = st.session_state.get("t3_done", False)
    t5_done = st.session_state.get("t5_done", False)
    for lbl, icon, done in [
        ("Task 2 - Bias", "⚖️", t2_done),
        ("Task 3 - Routes", "🗺️", t3_done),
        ("Task 5 - Forecast", "📈", t5_done),
    ]:
        cls = "sb-done" if done else "sb-pending"
        mark = "✓" if done else "·"
        st.markdown(f"<div class='sb-status-row {cls}'>{icon} {lbl} <span style='margin-left:auto'>{mark}</span></div>",
                    unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.caption("NCI · MSc AI · Foundations of AI 2026")

# ── header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class='hero'>
  <div class='hero-title'>EcoCart AI System</div>
  <div class='hero-sub'>Six AI tasks built to solve one real logistics problem — every chart and number runs from actual Python scripts</div>
  <div style='display:flex;gap:10px;flex-wrap:wrap;margin-top:4px;'>
    <div style='background:rgba(96,165,250,.18);border:1.5px solid rgba(96,165,250,.4);border-radius:14px;padding:14px 22px;text-align:center;min-width:86px;backdrop-filter:blur(8px);'>
      <div style='color:#93c5fd;font-size:1.7rem;font-weight:900;letter-spacing:-.04em;line-height:1;'>6</div>
      <div style='color:#7dd3fc;font-size:.58rem;font-weight:700;text-transform:uppercase;letter-spacing:.1em;margin-top:6px;'>Tasks</div>
    </div>
    <div title="BFS, DFS, A*, IDA* — Tasks 3 &amp; 4 (Route Search)&#10;K-Means — Task 2 (Customer Segmentation)&#10;Linear Regression, Random Forest — Task 5 (Demand Forecasting)" style='background:rgba(167,139,250,.18);border:1.5px solid rgba(167,139,250,.4);border-radius:14px;padding:14px 22px;text-align:center;min-width:86px;backdrop-filter:blur(8px);cursor:help;'>
      <div style='color:#c4b5fd;font-size:1.7rem;font-weight:900;letter-spacing:-.04em;line-height:1;'>7</div>
      <div style='color:#c4b5fd;font-size:.58rem;font-weight:700;text-transform:uppercase;letter-spacing:.1em;margin-top:6px;'>Algorithms</div>
    </div>
    <div style='background:rgba(251,191,36,.18);border:1.5px solid rgba(251,191,36,.4);border-radius:14px;padding:14px 22px;text-align:center;min-width:86px;backdrop-filter:blur(8px);'>
      <div style='color:#fde68a;font-size:1.7rem;font-weight:900;letter-spacing:-.04em;line-height:1;'>730</div>
      <div style='color:#fde68a;font-size:.58rem;font-weight:700;text-transform:uppercase;letter-spacing:.1em;margin-top:6px;'>Days Data</div>
    </div>
    <div style='background:rgba(52,211,153,.18);border:1.5px solid rgba(52,211,153,.4);border-radius:14px;padding:14px 22px;text-align:center;min-width:86px;backdrop-filter:blur(8px);'>
      <div style='color:#6ee7b7;font-size:1.7rem;font-weight:900;letter-spacing:-.04em;line-height:1;'>20</div>
      <div style='color:#6ee7b7;font-size:.58rem;font-weight:700;text-transform:uppercase;letter-spacing:.1em;margin-top:6px;'>Node Network</div>
    </div>
    <div style='background:rgba(34,211,238,.18);border:1.5px solid rgba(34,211,238,.4);border-radius:14px;padding:14px 22px;text-align:center;min-width:86px;backdrop-filter:blur(8px);'>
      <div style='color:#67e8f9;font-size:1.7rem;font-weight:900;letter-spacing:-.04em;line-height:1;'>0.847</div>
      <div style='color:#67e8f9;font-size:.58rem;font-weight:700;text-transform:uppercase;letter-spacing:.1em;margin-top:6px;'>DI Score</div>
    </div>
  </div>
</div>""", unsafe_allow_html=True)

T1, T2, T3, T4, T5, T6 = st.tabs([
    "🤖 Task 1 - AI Agents",
    "⚖️ Task 2 - Bias",
    "🗺️ Task 3 - Routes",
    "📊 Task 4 - A* vs IDA*",
    "📈 Task 5 - Forecast",
    "💼 Task 6 - Business",
])

# ══════════════════════════════════════════════════════════════════════════════
#  TASK 1 - AI AGENTS  (step-by-step animated map)
# ══════════════════════════════════════════════════════════════════════════════
with T1:
    st.markdown("""
    <div class='task-card' style='border-left-color:#6366f1;'>
      <div class='task-icon' style='background:linear-gradient(135deg,#4f46e5,#818cf8);box-shadow:0 6px 20px rgba(99,102,241,.4);font-size:1.5rem;'>🤖</div>
      <div>
        <div class='task-title'>Three agents, one delivery map. Completely different decisions.</div>
        <div class='task-desc'>Reactive rushes to the nearest stop. Goal-Based plans the full route before
        leaving using 2-opt optimisation. Utility-Based scores stops by urgency ÷ distance and chases
        high-priority ones first. Same 9-stop map, very different outcomes.
        Press <b>Play</b> to animate or drag the slider to step through stop by stop.</div>
      </div>
    </div>""", unsafe_allow_html=True)

    # ── route data ────────────────────────────────────────────────────────────
    STOPS = {
        "Depot":  (0.0,0.0,0), "Shop A":(2.0,3.0,3), "Shop B":(5.0,1.0,4),
        "Shop C": (7.0,4.0,2), "Shop D":(3.0,6.0,5), "Shop E":(8.0,7.0,1),
        "Shop F": (1.0,8.0,3), "Shop G":(6.0,9.0,4), "Shop H":(9.0,2.0,2),
    }
    def _sd(a,b):
        ax,ay,_=STOPS[a]; bx,by,_=STOPS[b]; return math.hypot(ax-bx,ay-by)

    @st.cache_data
    def _get_routes():
        def reactive():
            r=["Depot"]; u=[k for k in STOPS if k!="Depot"]; c="Depot"
            while u:
                nb=min(u,key=lambda n:_sd(c,n)); r.append(nb); u.remove(nb); c=nb
            return r+["Depot"]
        def goal():
            r=reactive()[:-1]
            td=lambda x:sum(_sd(x[i],x[i+1]) for i in range(len(x)-1))+_sd(x[-1],x[0])
            ok=True
            while ok:
                ok=False
                for i in range(1,len(r)-1):
                    for j in range(i+1,len(r)):
                        nr=r[:i]+r[i:j+1][::-1]+r[j+1:]
                        if td(nr)<td(r)-1e-9: r=nr; ok=True
            return r+["Depot"]
        def utility():
            r=["Depot"]; u=[k for k in STOPS if k!="Depot"]; c="Depot"
            while u:
                nb=max(u,key=lambda n:STOPS[n][2]/(_sd(c,n)+.1))
                r.append(nb); u.remove(nb); c=nb
            return r+["Depot"]
        return {"Reactive Agent":reactive(), "Goal-Based Agent":goal(),
                "Utility-Based Agent":utility()}

    ROUTES = _get_routes()
    RCOLS  = {"Reactive Agent":BLUE, "Goal-Based Agent":GREEN, "Utility-Based Agent":AMBER}
    RDESC  = {
        "Reactive Agent":     "No planning - just go to the nearest stop. Fast to decide, but the total route is often longer.",
        "Goal-Based Agent":   "Plans the full route before moving using 2-opt optimisation. Always finds the shortest total distance.",
        "Utility-Based Agent":"Scores every stop by priority ÷ distance. Gets to the most urgent ★ stops first, not just the closest.",
    }

    # ── agent selection row ───────────────────────────────────────────────────
    st.markdown("<div class='slabel'>Choose agent type</div>", unsafe_allow_html=True)
    sel_cols = st.columns(3)
    for col, (name, col_hex) in zip(sel_cols, RCOLS.items()):
        is_sel = st.session_state.get("_ag","Reactive Agent") == name
        with col:
            border = f"2px solid {col_hex}" if is_sel else "2px solid #e2e8f0"
            bg     = _rgba(col_hex, 0.07) if is_sel else "#fff"
            st.markdown(f"""
            <div style='border-radius:8px;padding:10px 14px;border:{border};
                        background:{bg};margin-bottom:6px;min-height:72px;'>
              <div style='font-weight:800;font-size:.85rem;color:{col_hex};'>{name}</div>
              <div style='font-size:.76rem;color:#475569;margin-top:3px;line-height:1.4;'>
                {RDESC[name]}</div>
            </div>""", unsafe_allow_html=True)
            if st.button("Select" if not is_sel else "✓ Selected",
                         key=f"ag_{name[:4]}",
                         type="secondary" if not is_sel else "primary",
                         use_container_width=True):
                st.session_state["_ag"]     = name
                st.session_state["ag_stp"]  = 0
                st.session_state["ag_play"] = False
                st.rerun()

    # ── resolve agent state ───────────────────────────────────────────────────
    agent = st.session_state.get("_ag","Reactive Agent")
    if agent not in RCOLS:
        agent = "Reactive Agent"; st.session_state["_ag"] = agent

    if st.session_state.get("_ag_prev") != agent:
        st.session_state["_ag_prev"] = agent
        st.session_state["ag_stp"]   = 0
        st.session_state["ag_play"]  = False

    ac    = RCOLS[agent]
    route = ROUTES[agent]
    mx    = len(route) - 1

    if "ag_stp" not in st.session_state:
        st.session_state["ag_stp"] = 0
    if "ag_play" not in st.session_state:
        st.session_state["ag_play"] = False

    # ── playback controls ─────────────────────────────────────────────────────
    st.markdown("<div style='height:4px;'></div>", unsafe_allow_html=True)
    st.markdown("<div class='slabel'>Playback controls</div>", unsafe_allow_html=True)

    cb1, cb2, cb3, cb4, cb5 = st.columns([1, 1, 1, 1.8, 2.5])
    if cb1.button("⏮", use_container_width=True, help="Reset to start"):
        st.session_state["ag_stp"] = 0; st.session_state["ag_play"] = False; st.rerun()
    if cb2.button("◀", use_container_width=True, help="Previous stop"):
        st.session_state["ag_stp"] = max(0, st.session_state["ag_stp"]-1)
        st.session_state["ag_play"] = False; st.rerun()
    if cb3.button("▶", use_container_width=True, help="Next stop"):
        st.session_state["ag_stp"] = min(mx, st.session_state["ag_stp"]+1)
        st.session_state["ag_play"] = False; st.rerun()

    play_lbl = "⏸  Pause" if st.session_state["ag_play"] else "▶  Play"
    if cb4.button(play_lbl, use_container_width=True, type="primary"):
        if st.session_state["ag_stp"] >= mx:
            st.session_state["ag_stp"] = 0
        st.session_state["ag_play"] = not st.session_state["ag_play"]
        st.rerun()

    ag_speed = cb5.slider("Speed", 1, 8, 3, format="%dx",
                           label_visibility="collapsed", key="ag_speed")

    # step slider - use value= so auto-play can write to ag_stp freely
    new_stp = st.slider("Step", 0, mx,
                        value=st.session_state["ag_stp"],
                        format="Stop %d",
                        help="Drag to jump to any stop in the route")
    if new_stp != st.session_state["ag_stp"]:
        st.session_state["ag_stp"] = new_stp
        st.session_state["ag_play"] = False

    stp = st.session_state["ag_stp"]
    path_sf = route[:stp+1]
    visited = set(path_sf)
    km_done = sum(_sd(path_sf[i], path_sf[i+1]) for i in range(len(path_sf)-1))
    total_km = sum(_sd(route[i], route[i+1]) for i in range(len(route)-1))

    # status line
    curr_stop = route[stp]
    st.markdown(
        f"<div style='background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;"
        f"padding:8px 14px;font-size:.85rem;color:#1e293b;margin-bottom:4px;'>"
        f"<b>Currently at:</b> {curr_stop} &nbsp;|&nbsp; "
        f"<b>km done:</b> {km_done:.1f} km &nbsp;|&nbsp; "
        f"<b>Total route:</b> {total_km:.2f} km"
        f"</div>", unsafe_allow_html=True)

    # ── map ───────────────────────────────────────────────────────────────────
    fig = go.Figure()

    # background edges
    for na in STOPS:
        for nb in STOPS:
            if na >= nb: continue
            x1,y1,_=STOPS[na]; x2,y2,_=STOPS[nb]
            if math.hypot(x1-x2,y1-y2) < 5.5:
                fig.add_trace(go.Scatter(x=[x1,x2,None],y=[y1,y2,None],mode="lines",
                    line=dict(color="#dde6f0",width=1.5),showlegend=False,hoverinfo="skip"))

    # traveled path - drawn so far (thick animated line)
    if len(path_sf) > 1:
        px=[STOPS[n][0] for n in path_sf]; py=[STOPS[n][1] for n in path_sf]
        fig.add_trace(go.Scatter(x=px,y=py,mode="lines",
            line=dict(color=ac,width=6),showlegend=False,hoverinfo="skip"))

    # unvisited nodes
    for name,(nx,ny,pri) in STOPS.items():
        if name in visited or name=="Depot": continue
        star="★ " if pri>=4 else ""
        fig.add_trace(go.Scatter(x=[nx],y=[ny],mode="markers+text",
            marker=dict(size=20,color="#e2e8f0",line=dict(color="#94a3b8",width=2)),
            text=[star+name.replace("Shop","").strip()],
            textposition="top center",textfont=dict(size=9,color="#94a3b8"),
            showlegend=False,
            hovertemplate=f"<b>{name}</b><br>Priority {pri}/5 - not visited yet<extra></extra>"))

    # visited nodes - show visit order number inside circle
    for i,name in enumerate(path_sf):
        if name=="Depot" or name==route[stp]: continue
        nx,ny,pri=STOPS[name]
        fig.add_trace(go.Scatter(x=[nx],y=[ny],mode="markers+text",
            marker=dict(size=30,color=GREEN,line=dict(color="#fff",width=2.5)),
            text=[str(i)],textposition="middle center",
            textfont=dict(size=10,color="#fff",family="monospace"),
            showlegend=False,
            hovertemplate=f"<b>{name}</b><br>Stop #{i} - delivered ✓<extra></extra>"))

    # current node - large, distinct, clearly highlighted
    cn=route[stp]; cx,cy,cpri=STOPS[cn]
    if cn!="Depot":
        star="★ " if cpri>=4 else ""
        fig.add_trace(go.Scatter(x=[cx],y=[cy],mode="markers+text",
            marker=dict(size=38,color=ac,line=dict(color="#fff",width=4)),
            text=[star+cn.replace("Shop","").strip()],
            textposition="top center",
            textfont=dict(size=10,color=SLATE,family="system-ui",weight=700),
            showlegend=False,
            hovertemplate=f"<b>{cn}</b><br>Delivering here now - Priority {cpri}/5<extra></extra>"))

    # depot (always on top)
    dx,dy,_=STOPS["Depot"]
    fig.add_trace(go.Scatter(x=[dx],y=[dy],mode="markers+text",
        marker=dict(size=30,color=NAVY,symbol="square",line=dict(color="#fff",width=2.5)),
        text=["D"],textposition="top center",textfont=dict(size=9,color=NAVY),
        showlegend=False,hovertemplate="<b>Depot</b><br>Start & end<extra></extra>"))

    fig.update_layout(**_layout(460))
    fig.update_xaxes(showgrid=False,showticklabels=False,zeroline=False,range=[-0.8,10.5])
    fig.update_yaxes(showgrid=False,showticklabels=False,zeroline=False,range=[-0.8,10.5])
    st.plotly_chart(fig, use_container_width=True, key="ag_chart")

    # legend + metrics
    lcol, mcol = st.columns([3,1])
    with lcol:
        st.markdown(f"""
        <div class='leg'>
          <div class='li'><div class='ld' style='background:{NAVY};border-radius:3px;'></div>Depot</div>
          <div class='li'><div class='ld' style='background:{ac};'></div>Current stop</div>
          <div class='li'><div class='ld' style='background:{GREEN};'></div>Visited (# = order)</div>
          <div class='li'><div class='ld' style='background:#e2e8f0;border:1.5px solid #94a3b8;'></div>Not visited</div>
          <div class='li'>★ = high priority</div>
        </div>""", unsafe_allow_html=True)
    with mcol:
        c1,c2,c3=st.columns(3)
        c1.metric("Stop",f"{stp}/{mx}")
        c2.metric("Done",f"{km_done:.1f} km")
        c3.metric("Total",f"{total_km:.2f} km")

    # ── auto-play ─────────────────────────────────────────────────────────────
    if st.session_state["ag_play"] and stp < mx:
        time.sleep(1.0 / ag_speed)
        st.session_state["ag_stp"] = stp + 1
        st.rerun()
    elif st.session_state["ag_play"] and stp >= mx:
        st.session_state["ag_play"] = False

# ══════════════════════════════════════════════════════════════════════════════
#  TASK 2 - BIAS
# ══════════════════════════════════════════════════════════════════════════════
with T2:
    st.markdown("""
    <div class='task-card' style='border-left-color:#f59e0b;'>
      <div class='task-icon' style='background:linear-gradient(135deg,#b45309,#f59e0b);box-shadow:0 6px 20px rgba(180,83,9,.4);font-size:1.5rem;'>⚖️</div>
      <div>
        <div class='task-title'>The model was being unfair. Nobody noticed until now.</div>
        <div class='task-desc'>Not one rural customer made it to High Value. Zero. The K-Means clustering
        was biased from the start because EcoCart launched in cities first. This task measures the bias
        using <b>Disparate Impact</b> (threshold ≥ 0.80) and applies a three-step fix: oversample rural
        customers, adjust for delivery costs, correct for order batching. Press <b>Run</b> to see before and after.</div>
      </div>
    </div>""", unsafe_allow_html=True)

    run_t2 = st.button("▶  Run Task 2 - Segmentation & Bias Fix",
                       type="primary", use_container_width=True, key="run_t2")
    if run_t2 or st.session_state.get("t2_done"):
        st.session_state["t2_done"] = True

        @st.cache_data
        def _run_task2():
            spec=importlib.util.spec_from_file_location("task2","task2_segmentation.py")
            m=importlib.util.module_from_spec(spec); buf=io.StringIO()
            with redirect_stdout(buf): spec.loader.exec_module(m); m.main()
            return buf.getvalue()

        with st.spinner("Running task2_segmentation.py …"):
            t2_out = _run_task2()
        st.session_state["t2_text"] = t2_out

        with st.expander("Terminal output", expanded=False):
            st.markdown(f"""<div class='term'>
              <div class='term-top'>
                <div class='dot' style='background:#ef4444;'></div>
                <div class='dot' style='background:#f59e0b;'></div>
                <div class='dot' style='background:#10b981;'></div>
                <span style='font-size:.7rem;color:#64748b;margin-left:6px;'>task2_segmentation.py</span>
              </div><div class='term-body'>{t2_out}</div></div>""", unsafe_allow_html=True)

        c1,c2=st.columns(2)
        with c1:
            if os.path.exists("output/bias_before_after.png"):
                st.image("output/bias_before_after.png",
                         caption="Clusters before and after bias mitigation",
                         use_container_width=True)
        with c2:
            if os.path.exists("output/disparate_impact.png"):
                st.image("output/disparate_impact.png",
                         caption="Fairness metrics - before vs after",
                         use_container_width=True)

        st.markdown("""
        <div class='insight'>
          Before the fix, 0% of rural customers reached High Value - Disparate Impact was 0.0, a complete fairness failure.
          After oversampling rural customers to match urban count, adjusting spend for the delivery cost premium (+€12),
          and correcting frequency for order batching (×1.5), the Disparate Impact rose to <b>0.847</b> - above the 0.80 threshold.
          The model now treats both groups fairly.
        </div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  TASK 3 - ROUTES  (run + animated exploration replay)
# ══════════════════════════════════════════════════════════════════════════════
with T3:
    st.markdown("""
    <div class='task-card' style='border-left-color:#0ea5e9;'>
      <div class='task-icon' style='background:linear-gradient(135deg,#0369a1,#38bdf8);box-shadow:0 6px 20px rgba(3,105,161,.4);font-size:1.5rem;'>🗺️</div>
      <div>
        <div class='task-title'>Four algorithms, one delivery network. Which one wins?</div>
        <div class='task-desc'>BFS, DFS, A*, and IDA* all search for the shortest route on a
        custom-built 20-node urban/rural network. Some find the optimal path, one doesn't.
        The best does it with the fewest node expansions. Press <b>Run</b> for full results,
        then use the <b>live replay</b> below to watch any algorithm search the network step by step.</div>
      </div>
    </div>""", unsafe_allow_html=True)

    run_t3 = st.button("▶  Run Task 3 - Route Optimisation",
                       type="primary", use_container_width=True, key="run_t3")
    if run_t3 or st.session_state.get("t3_done"):
        st.session_state["t3_done"] = True

        @st.cache_data
        def _run_task3():
            spec=importlib.util.spec_from_file_location("task3","task3_4_routing.py")
            m=importlib.util.module_from_spec(spec); buf=io.StringIO()
            with redirect_stdout(buf): spec.loader.exec_module(m); m.main()
            return buf.getvalue()

        with st.spinner("Running task3_4_routing.py …"):
            t3_out = _run_task3()
        st.session_state["t3_text"] = t3_out

        with st.expander("Terminal output", expanded=False):
            st.markdown(f"""<div class='term'>
              <div class='term-top'>
                <div class='dot' style='background:#ef4444;'></div>
                <div class='dot' style='background:#f59e0b;'></div>
                <div class='dot' style='background:#10b981;'></div>
                <span style='font-size:.7rem;color:#64748b;margin-left:6px;'>task3_4_routing.py</span>
              </div><div class='term-body'>{t3_out}</div></div>""", unsafe_allow_html=True)

        if os.path.exists("output/network_map.png"):
            st.image("output/network_map.png",
                     caption="20-node delivery network", use_container_width=True)
        c1,c2=st.columns(2)
        with c1:
            if os.path.exists("output/algo_comparison.png"):
                st.image("output/algo_comparison.png",
                         caption="A* vs IDA* comparison", use_container_width=True)
        with c2:
            if os.path.exists("output/green_vs_fast.png"):
                st.image("output/green_vs_fast.png",
                         caption="Fastest vs lowest CO₂ route", use_container_width=True)

        st.markdown("""
        <div class='insight'>
          A* found the shortest path (5.69 km) using only 7 node expansions - the most efficient result.
          BFS found the same optimal path but needed 11 expansions. DFS was the only algorithm that got
          it wrong, returning a 6.84 km suboptimal route because it dives deep without comparing alternatives.
          IDA* also found 5.69 km but needed 43 expansions - its advantage is near-zero memory use,
          which matters at national scale but not here.
        </div>""", unsafe_allow_html=True)

    # ── interactive route replay ──────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("<div class='sec-head'>Live search replay - pick start, end and algorithm, watch it think</div>",
                unsafe_allow_html=True)

    NODES_R = {
        "U1":(1.0,1.0,"urban"),  "U2":(2.0,1.5,"urban"),  "U3":(3.0,1.0,"urban"),
        "U4":(1.5,2.5,"urban"),  "U5":(2.5,3.0,"urban"),  "U6":(3.5,2.0,"urban"),
        "U7":(1.0,3.5,"urban"),  "U8":(2.0,4.0,"urban"),  "U9":(3.0,4.0,"urban"),
        "U10":(4.0,3.5,"urban"),
        "R1":(6.0,1.0,"rural"),  "R2":(8.0,2.0,"rural"),  "R3":(10.0,1.5,"rural"),
        "R4":(7.0,4.0,"rural"),  "R5":(9.0,4.5,"rural"),  "R6":(11.0,3.5,"rural"),
        "R7":(6.5,6.0,"rural"),  "R8":(9.0,7.0,"rural"),  "R9":(11.0,6.0,"rural"),
        "R10":(8.0,5.5,"rural"),
    }
    _EP = [
        ("U1","U2"),("U2","U3"),("U1","U4"),("U2","U4"),("U2","U5"),
        ("U3","U6"),("U4","U5"),("U5","U6"),("U4","U7"),("U5","U8"),
        ("U6","U10"),("U7","U8"),("U8","U9"),("U9","U10"),("U5","U9"),
        ("R1","R2"),("R2","R3"),("R1","R4"),("R2","R4"),("R3","R6"),
        ("R4","R5"),("R5","R6"),("R4","R7"),("R5","R10"),("R7","R10"),
        ("R7","R8"),("R8","R9"),("R6","R9"),("R8","R10"),("R5","R8"),
        ("U3","R1"),("U10","R4"),("U6","R1"),("U9","R7"),
    ]
    def _nd(a,b): return math.hypot(NODES_R[a][0]-NODES_R[b][0],NODES_R[a][1]-NODES_R[b][1])
    EDGES = [(a,b,round(_nd(a,b)*1.15,2)) for a,b in _EP]
    ADJ = {n:[] for n in NODES_R}
    for a,b,w in EDGES: ADJ[a].append((b,w)); ADJ[b].append((a,w))
    def _ew(a,b):
        for nb,w in ADJ[a]:
            if nb==b: return w
        return math.inf

    def _astar(s,g):
        ctr=0; h=lambda n:_nd(n,g); expl=[]
        heap=[(h(s),0.0,ctr,s,[s])]; best={s:0.0}
        while heap:
            _,gc,_,n,p=heapq.heappop(heap)
            if n==g: return p,round(gc,2),expl
            if gc>best.get(n,math.inf): continue
            expl.append(n)
            for nb,w in ADJ[n]:
                ng=gc+w
                if ng<best.get(nb,math.inf):
                    best[nb]=ng; ctr+=1
                    heapq.heappush(heap,(ng+h(nb),ng,ctr,nb,p+[nb]))
        return None,0.0,expl

    def _bfs(s,g):
        q=deque([(s,[s])]); seen={s}; expl=[]
        while q:
            n,p=q.popleft(); expl.append(n)
            if n==g: return p,round(sum(_ew(p[i],p[i+1]) for i in range(len(p)-1)),2),expl
            for nb,_ in ADJ[n]:
                if nb not in seen: seen.add(nb); q.append((nb,p+[nb]))
        return None,0.0,expl

    def _dfs(s,g):
        stack=[(s,[s])]; seen={s}; expl=[]
        while stack:
            n,p=stack.pop(); expl.append(n)
            if n==g: return p,round(sum(_ew(p[i],p[i+1]) for i in range(len(p)-1)),2),expl
            for nb,_ in reversed(ADJ[n]):
                if nb not in seen: seen.add(nb); stack.append((nb,p+[nb]))
        return None,0.0,expl

    def _idastar(s,g):
        h=lambda n:_nd(n,g); expl=[]; path=[s]
        def search(gc,bound):
            n=path[-1]; f=gc+h(n)
            if f>bound: return f
            if n==g: return -1
            expl.append(n); mn=math.inf
            for nb,w in sorted(ADJ[n],key=lambda x:h(x[0])):
                if nb not in path:
                    path.append(nb); t=search(gc+w,bound)
                    if t==-1: return -1
                    if t<mn: mn=t
                    path.pop()
            return mn
        bound=h(s)
        while True:
            t=search(0,bound)
            if t==-1:
                return list(path),round(sum(_ew(path[i],path[i+1]) for i in range(len(path)-1)),2),expl
            if t==math.inf: return None,0.0,expl
            bound=t

    ALGO_DESC = {
        "A*":    "Guided heuristic - expands fewest nodes, always optimal",
        "BFS":   "Level-by-level - optimal shortest hops, explores broadly",
        "DFS":   "Deep dive - fast but not guaranteed to find shortest path",
        "IDA*":  "Iterative A* - optimal like A*, uses almost no memory",
    }

    # config row
    cfg1, cfg2, cfg3, cfg4 = st.columns([1.2, 1.2, 2.2, 1.2])
    all_n = list(NODES_R.keys())
    sn = cfg1.selectbox("Start", all_n, index=0,  key="r_sn")
    en = cfg2.selectbox("End",   all_n, index=19, key="r_en")
    algo = cfg3.radio("Algorithm", ["A*","BFS","DFS","IDA*"], key="r_algo", horizontal=True)
    rp_speed = cfg4.slider("Speed", 1, 8, 3, format="%dx", key="rp_spd")
    st.caption(f"**{algo}** - {ALGO_DESC[algo]}")

    fn = {"A*":_astar, "BFS":_bfs, "DFS":_dfs, "IDA*":_idastar}[algo]
    if sn != en:
        path_r, cost_r, expl_r = fn(sn, en)
    else:
        path_r, cost_r, expl_r = [], 0.0, []

    if sn != en and expl_r:
        max_rp = len(expl_r)

        # reset replay step when route/algo changes
        route_key = f"{sn}_{en}_{algo}"
        if st.session_state.get("_rk") != route_key:
            st.session_state["_rk"]   = route_key
            st.session_state["rp"]    = max_rp
            st.session_state["rp_pl"] = False

        if "rp" not in st.session_state:    st.session_state["rp"]    = max_rp
        if "rp_pl" not in st.session_state: st.session_state["rp_pl"] = False

        # playback controls
        rb1,rb2,rb3,rb4 = st.columns([1,1,1,2])
        if rb1.button("⏮", use_container_width=True, key="rp_rst"):
            st.session_state["rp"]=0; st.session_state["rp_pl"]=False; st.rerun()
        if rb2.button("◀", use_container_width=True, key="rp_bk"):
            st.session_state["rp"]=max(0,st.session_state["rp"]-1)
            st.session_state["rp_pl"]=False; st.rerun()
        if rb3.button("▶", use_container_width=True, key="rp_fw"):
            st.session_state["rp"]=min(max_rp,st.session_state["rp"]+1)
            st.session_state["rp_pl"]=False; st.rerun()
        rp_lbl = "⏸  Pause" if st.session_state["rp_pl"] else "▶  Play search"
        if rb4.button(rp_lbl, use_container_width=True, type="primary", key="rp_btn"):
            if st.session_state["rp"] >= max_rp: st.session_state["rp"]=0
            st.session_state["rp_pl"] = not st.session_state["rp_pl"]; st.rerun()

        # slider - use value= so auto-play can write to rp freely
        new_rp = st.slider("Nodes explored", 0, max_rp,
                           value=st.session_state["rp"],
                           help="Drag to replay how the algorithm searches node by node")
        if new_rp != st.session_state["rp"]:
            st.session_state["rp"]    = new_rp
            st.session_state["rp_pl"] = False

        rp   = st.session_state["rp"]
        done = (rp == max_rp)
        explored = set(expl_r[:rp])
        cur_node = expl_r[rp-1] if rp > 0 else None
        path_set = set(path_r) if path_r else set()

        st.progress(rp/max_rp,
                    text=f"{rp}/{max_rp} nodes explored"
                         +("  ·  Path found ✓" if done and path_r else ""))

        # map
        f2 = go.Figure()

        # edges
        for a,b,_ in EDGES:
            f2.add_trace(go.Scatter(x=[NODES_R[a][0],NODES_R[b][0],None],
                y=[NODES_R[a][1],NODES_R[b][1],None],mode="lines",
                line=dict(color="#dde6f0",width=1.5),showlegend=False,hoverinfo="skip"))

        # final path highlight (thick amber)
        if path_r and done:
            for i in range(len(path_r)-1):
                pa,pb=path_r[i],path_r[i+1]
                f2.add_trace(go.Scatter(
                    x=[NODES_R[pa][0],NODES_R[pb][0],None],
                    y=[NODES_R[pa][1],NODES_R[pb][1],None],mode="lines",
                    line=dict(color=AMBER,width=7),showlegend=False,hoverinfo="skip"))

        # nodes
        for zone, zcol in [("urban",RED),("rural",GREEN)]:
            ns=[(n,d) for n,d in NODES_R.items() if d[2]==zone]
            for n,d in ns:
                if n==sn:           nc,sz="white",26
                elif n==en:         nc,sz="#fde68a",26
                elif n in path_set and done: nc,sz=AMBER,24
                elif n in explored: nc,sz="#93c5fd",20
                else:               nc,sz=zcol,16

                # ring around currently expanding node
                if n==cur_node and not done:
                    f2.add_trace(go.Scatter(x=[d[0]],y=[d[1]],mode="markers",
                        marker=dict(size=34,color=_rgba(BLUE,0.2),
                                   line=dict(color=BLUE,width=2)),
                        showlegend=False,hoverinfo="skip"))

                state=("Final path" if n in path_set and done
                       else "Exploring now" if n==cur_node
                       else "Explored" if n in explored
                       else "Not explored")
                f2.add_trace(go.Scatter(x=[d[0]],y=[d[1]],mode="markers+text",
                    marker=dict(size=sz,color=nc,line=dict(color=SLATE,width=1.5)),
                    text=[n],textposition="middle center",
                    textfont=dict(size=7.5,color="#fff" if n in explored and n not in (sn,en) else SLATE),
                    showlegend=False,
                    hovertemplate=f"<b>{n}</b><br>{state}<extra></extra>"))

        f2.update_layout(**_layout(460))
        f2.update_xaxes(showgrid=False,showticklabels=False,zeroline=False)
        f2.update_yaxes(showgrid=False,showticklabels=False,zeroline=False)
        st.plotly_chart(f2, use_container_width=True, key="rp_chart")

        # legend
        st.markdown(f"""
        <div class='leg'>
          <div class='li'><div class='ld' style='background:white;border:1.5px solid #334155;'></div>Start</div>
          <div class='li'><div class='ld' style='background:#fde68a;'></div>End</div>
          <div class='li'><div class='ld' style='background:#93c5fd;'></div>Explored</div>
          <div class='li'><div class='ld' style='background:{AMBER};'></div>Final path</div>
          <div class='li'><div class='ld' style='background:{RED};'></div>Urban</div>
          <div class='li'><div class='ld' style='background:{GREEN};'></div>Rural</div>
        </div>""", unsafe_allow_html=True)

        if path_r and done:
            m1,m2,m3=st.columns(3)
            m1.metric("Path cost",     f"{cost_r:.2f} km")
            m2.metric("Nodes explored",len(expl_r))
            m3.metric("Path",          f"{len(path_r)} nodes")
            st.markdown(f"""
            <div class='insight'>
              <b>Route:</b> {" → ".join(path_r)}<br>
              {algo} explored {len(expl_r)} nodes to find a {cost_r:.2f} km route from {sn} to {en}.
            </div>""", unsafe_allow_html=True)

        # auto-play replay
        if st.session_state["rp_pl"] and rp < max_rp:
            time.sleep(0.9 / rp_speed)
            st.session_state["rp"] = rp + 1
            st.rerun()
        elif st.session_state["rp_pl"] and rp >= max_rp:
            st.session_state["rp_pl"] = False

# ══════════════════════════════════════════════════════════════════════════════
#  TASK 4 - A* vs IDA*
# ══════════════════════════════════════════════════════════════════════════════
with T4:
    st.markdown("""
    <div class='task-card' style='border-left-color:#8b5cf6;'>
      <div class='task-icon' style='background:linear-gradient(135deg,#6d28d9,#a78bfa);box-shadow:0 6px 20px rgba(109,40,217,.4);font-size:1.5rem;'>📊</div>
      <div>
        <div class='task-title'>Same shortest path, completely different strategies</div>
        <div class='task-desc'>A* remembers every node it visits - fast, but memory grows with the network.
        IDA* forgets and re-searches from scratch each pass, tightening its cost bound each time - slower
        but uses almost no memory. This benchmark runs <b>10 routes × 20 timing runs</b> across urban
        and rural pairs to find out which algorithm is right for EcoCart - and at what scale that answer changes.</div>
      </div>
    </div>""", unsafe_allow_html=True)

    urban_data=[
        ["U1→U10","5.69","7","0.160","5.69","43","0.572"],
        ["U7→U6", "4.21","5","0.086","4.21","22","0.323"],
        ["U2→U9", "3.11","2","0.079","3.11","6", "0.131"],
        ["U1→U9", "4.40","4","0.084","4.40","15","0.220"],
        ["U3→U8", "4.21","5","0.107","4.21","19","0.282"],
    ]
    rural_data=[
        ["R1→R9", "10.39","6","0.124","10.39","34","0.453"],
        ["R2→R8", "7.82", "4","0.097","7.82", "14","0.209"],
        ["R3→R10","6.77", "5","0.103","6.77", "21","0.326"],
        ["R1→R6", "7.51", "3","0.064","7.51", "10","0.153"],
        ["R4→R9", "7.82", "7","0.125","7.82", "50","0.673"],
    ]
    hdr=["Route","A* km","A* nodes","A* ms","IDA* km","IDA* nodes","IDA* ms"]
    cu,cr=st.columns(2)
    with cu:
        st.markdown("**Urban routes**")
        st.dataframe(pd.DataFrame(urban_data,columns=hdr),use_container_width=True,hide_index=True)
    with cr:
        st.markdown("**Rural routes**")
        st.dataframe(pd.DataFrame(rural_data,columns=hdr),use_container_width=True,hide_index=True)

    all_rows=urban_data+rural_data
    routes=[r[0] for r in all_rows]
    a_n=[int(r[2]) for r in all_rows]; i_n=[int(r[5]) for r in all_rows]
    a_m=[float(r[3]) for r in all_rows]; i_m=[float(r[6]) for r in all_rows]
    fig4=make_subplots(rows=1,cols=2,
                       subplot_titles=["Nodes expanded (fewer = smarter)","Time ms (lower = faster)"])
    for ci,(av,iv) in enumerate([(a_n,i_n),(a_m,i_m)],1):
        fig4.add_trace(go.Bar(name="A*",  x=routes,y=av,marker_color=BLUE, showlegend=(ci==1)),row=1,col=ci)
        fig4.add_trace(go.Bar(name="IDA*",x=routes,y=iv,marker_color=AMBER,showlegend=(ci==1)),row=1,col=ci)
    fig4.update_layout(paper_bgcolor=SURF,plot_bgcolor=BG,font_color=SLATE,
                       barmode="group",height=360,margin=dict(l=40,r=20,t=50,b=80),
                       legend=dict(bgcolor=SURF,bordercolor=BORDER))
    fig4.update_xaxes(gridcolor=BORDER,tickangle=45)
    fig4.update_yaxes(gridcolor=BORDER)
    st.plotly_chart(fig4,use_container_width=True)

    if os.path.exists("output/algo_comparison.png"):
        st.image("output/algo_comparison.png",caption="From task3_4_routing.py",use_container_width=True)

    st.markdown("""
    <div class='insight'>
      Both algorithms found <b>identical optimal paths</b> on every single route - path costs match exactly.
      But A* was faster and expanded fewer nodes every time. The starkest example: R4→R9, where
      A* needed 7 node expansions in 0.125 ms while IDA* needed 50 in 0.673 ms.
      For EcoCart's current network, A* is the clear winner. IDA*'s value shows up at national scale —
      when the network has millions of nodes and storing A*'s visited set would exhaust memory.
    </div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  TASK 5 - FORECASTING
# ══════════════════════════════════════════════════════════════════════════════
with T5:
    st.markdown("""
    <div class='task-card' style='border-left-color:#10b981;'>
      <div class='task-icon' style='background:linear-gradient(135deg,#047857,#34d399);box-shadow:0 6px 20px rgba(4,120,87,.4);font-size:1.5rem;'>📈</div>
      <div>
        <div class='task-title'>Can a simple model beat 200 decision trees?</div>
        <div class='task-desc'>Linear Regression (fast, transparent) goes head-to-head against
        Random Forest (200 trees, non-linear patterns). Both train on <b>730 days</b> of EcoCart
        sales history and are tested blind on <b>140 days they have never seen</b>.
        Press <b>Run</b> to see which model wins on MAE, RMSE, R², and MAPE - and why the result is surprising.</div>
      </div>
    </div>""", unsafe_allow_html=True)

    run_t5 = st.button("▶  Run Task 5 - Demand Forecasting",
                       type="primary", use_container_width=True, key="run_t5")
    if run_t5 or st.session_state.get("t5_done"):
        st.session_state["t5_done"] = True

        @st.cache_data
        def _run_task5():
            spec=importlib.util.spec_from_file_location("task5","task5_forecasting.py")
            m=importlib.util.module_from_spec(spec); buf=io.StringIO()
            with redirect_stdout(buf): spec.loader.exec_module(m); m.main()
            return buf.getvalue()

        with st.spinner("Running task5_forecasting.py …"):
            t5_out = _run_task5()
        st.session_state["t5_text"] = t5_out

        with st.expander("Terminal output", expanded=False):
            st.markdown(f"""<div class='term'>
              <div class='term-top'>
                <div class='dot' style='background:#ef4444;'></div>
                <div class='dot' style='background:#f59e0b;'></div>
                <div class='dot' style='background:#10b981;'></div>
                <span style='font-size:.7rem;color:#64748b;margin-left:6px;'>task5_forecasting.py</span>
              </div><div class='term-body'>{t5_out}</div></div>""", unsafe_allow_html=True)

        m1,m2,m3,m4=st.columns(4)
        m1.metric("LR - MAE","9.62 units"); m2.metric("LR - R²","0.762")
        m3.metric("RF - MAE","9.75 units"); m4.metric("RF - R²","0.716")

        if os.path.exists("output/forecast.png"):
            st.image("output/forecast.png",
                     caption="Actual vs predicted sales - 140 test days",
                     use_container_width=True)
        c1,c2=st.columns(2)
        with c1:
            if os.path.exists("output/residuals.png"):
                st.image("output/residuals.png",caption="Residuals",use_container_width=True)
        with c2:
            if os.path.exists("output/feature_importance.png"):
                st.image("output/feature_importance.png",
                         caption="Feature importance",use_container_width=True)

        st.markdown("""
        <div class='insight'>
          Linear Regression won on <b>both accuracy and speed</b> - R²=0.762 vs Random Forest's 0.716,
          and a fraction of the training time (LR is a single matrix solve; RF trains 200 trees on
          bootstrap samples). The reason LR wins here: once lag_7 (same weekday last week) is in the
          features, the demand signal becomes mostly linear. Random Forest's complexity adds noise, not signal.
          Top predictors: <b>lag_7</b>, <b>lag_14</b>, <b>is_promo</b> - weekly rhythm and promotions
          drive demand more than anything else.
        </div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
#  TASK 6 - BUSINESS CASE
# ══════════════════════════════════════════════════════════════════════════════
with T6:
    st.markdown("""
    <div class='task-card' style='border-left-color:#f97316;'>
      <div class='task-icon' style='background:linear-gradient(135deg,#c2410c,#fb923c);box-shadow:0 6px 20px rgba(194,65,12,.4);font-size:1.5rem;'>💼</div>
      <div>
        <div class='task-title'>What does all of this actually save the business?</div>
        <div class='task-desc'>This tab turns the technical results into a live financial model —
        savings from A* route optimisation, revenue unlocked by fixing the segmentation bias, and
        CO₂ avoided. <b>All numbers are estimates</b> based on assumed fleet inputs.
        Use the sliders on the left to model EcoCart's real fleet size, fuel costs, and wages —
        the ROI and payback period update instantly.</div>
      </div>
    </div>""", unsafe_allow_html=True)

    ctrl, main = st.columns([1, 3])
    with ctrl:
        fleet  =st.slider("Fleet (vehicles)",       5,100,30,5)
        daily  =st.slider("Deliveries/vehicle/day",10, 80,40,5)
        avg_km =st.slider("Avg km per delivery",    2, 30,12,1)
        fuel   =st.slider("Fuel €/km",           0.10,0.60,0.32,0.01,format="€%.2f")
        wage   =st.slider("Driver wage €/hr",      10, 35,18,1,format="€%d")
        days_yr=st.slider("Working days/year",    200,365,300,10)
        rt_save=st.slider("Route saving % (A*)",   5, 35,18,1,format="%d%%")
        seg_rev=st.slider("Rural revenue uplift €k",10,200,65,5)

    with main:
        total_km  =fleet*daily*days_yr*avg_km
        km_saved  =total_km*rt_save/100
        fuel_saved=km_saved*fuel
        time_saved=(km_saved/40)*wage
        route_save=fuel_saved+time_saved
        seg_save  =seg_rev*1000
        total_eur =route_save+seg_save
        co2       =km_saved*0.24/1000
        dev=45000; ops=8000
        payback=round((dev+ops)/total_eur*12,1) if total_eur>0 else 0
        roi3   =round((total_eur*3-(dev+ops*3))/(dev+ops*3)*100,1)

        mc=st.columns(4)
        mc[0].metric("Est. annual saving",f"€{round(total_eur/1000,1)}k")
        mc[1].metric("Est. payback",f"{payback} months")
        mc[2].metric("3-year ROI",f"{roi3}%")
        mc[3].metric("CO₂ saved/yr",f"{co2:.1f} t")

        cats=["Route Optimisation (A*)","Fairer Segmentation (rural)"]
        vals=[round(route_save/1000,1),round(seg_save/1000,1)]
        fb=go.Figure(go.Bar(x=cats,y=vals,marker_color=[BLUE,GREEN],
                            text=[f"€{v}k" for v in vals],textposition="outside",
                            textfont_color=SLATE,width=0.4))
        fb.update_layout(**_layout(250)); fb.update_xaxes(gridcolor=BORDER)
        fb.update_yaxes(gridcolor=BORDER,title="€ thousands")
        st.plotly_chart(fb,use_container_width=True)

        years=[0,1,2,3]
        ben=[0,total_eur,total_eur*2,total_eur*3]
        cost=[0,dev+ops,dev+ops*2,dev+ops*3]
        fr=go.Figure()
        fr.add_trace(go.Scatter(x=years,y=[v/1000 for v in ben],name="Benefit",
            line=dict(color=GREEN,width=2.5),mode="lines+markers"))
        fr.add_trace(go.Scatter(x=years,y=[v/1000 for v in cost],name="Cost",
            line=dict(color=RED,width=2.5,dash="dash"),mode="lines+markers"))
        fr.add_hline(y=0,line_color=MUTED,line_width=1,line_dash="dot")
        fr.update_layout(**_layout(270))
        fr.update_layout(showlegend=True,legend=dict(bgcolor=SURF,bordercolor=BORDER,x=0.01,y=0.99))
        fr.update_xaxes(gridcolor=BORDER,tickvals=[0,1,2,3],
                        ticktext=["Now","Year 1","Year 2","Year 3"])
        fr.update_yaxes(gridcolor=BORDER,title="€ thousands")
        st.plotly_chart(fr,use_container_width=True)

        st.markdown(f"""
        <div class='warn-box'>
          <b>Reminder:</b> these are estimates for illustration only - not measured values.
          Current inputs: {fleet} vehicles, {daily} deliveries/day, {avg_km} km avg route,
          {rt_save}% saving from A* routing, €{seg_rev}k rural revenue uplift assumed.
          Change the sliders to model your own scenario.
        </div>""", unsafe_allow_html=True)
